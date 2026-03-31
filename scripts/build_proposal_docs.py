#!/usr/bin/env python3
"""Build LaTeX, PDF, and DOCX artifacts for the AI4MH proposal.

The script keeps one structured content definition and renders:
- LaTeX source
- PDF via reportlab
- DOCX via python-docx

Default outputs are repo-local. Callers can override paths explicitly.
"""

from __future__ import annotations

import argparse
import html
from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class ParagraphBlock:
    text: str


@dataclass(frozen=True)
class BulletListBlock:
    items: tuple[str, ...]


@dataclass(frozen=True)
class NumberedListBlock:
    items: tuple[str, ...]


@dataclass(frozen=True)
class EquationBlock:
    plain: str
    latex: str


@dataclass(frozen=True)
class QuoteBlock:
    text: str


@dataclass(frozen=True)
class TableBlock:
    headers: tuple[str, ...]
    rows: tuple[tuple[str, ...], ...]


@dataclass(frozen=True)
class LabelValueBlock:
    label: str
    value: str


@dataclass(frozen=True)
class FigurePromptBlock:
    title: str
    placement: str
    prompt: str


@dataclass(frozen=True)
class Subsection:
    title: str
    blocks: tuple[object, ...]


@dataclass(frozen=True)
class Section:
    title: str
    blocks: tuple[object, ...]


@dataclass(frozen=True)
class ProposalDocument:
    title: str
    author: str
    sections: tuple[Section, ...]


def build_document() -> ProposalDocument:
    return ProposalDocument(
        title="AI4MH: Credible Signal Detection for Mental Health Crisis Monitoring",
        author="Ujjwal Singh",
        sections=(
            Section(
                title="Abstract",
                blocks=(
                    ParagraphBlock(
                        "Early detection of mental health crises using social media is limited "
                        "by sparse data, noisy signals, and adversarial manipulation. Existing "
                        "systems optimize classification accuracy, which fails under deployment "
                        "conditions and produces unreliable escalation decisions."
                    ),
                    ParagraphBlock(
                        "This project proposes a credible signal detection framework that treats "
                        "crisis monitoring as an uncertainty-aware inference problem. Instead of "
                        "direct classification, the system estimates a Confidence-Weighted Crisis "
                        "Index (chi) using sample stability, temporal consistency, and spatial "
                        "corroboration."
                    ),
                    ParagraphBlock(
                        "A threat-filtering layer suppresses non-organic signals using lexical "
                        "entropy, velocity anomalies, and diffusion patterns. The system emits "
                        "confidence-calibrated escalation decisions with mandatory human review. "
                        "The target outcome is decision reliability, not raw model accuracy."
                    ),
                ),
            ),
            Section(
                title="Problem Statement",
                blocks=(
                    ParagraphBlock(
                        "Current crisis-monitoring systems behave as classification pipelines. "
                        "They fail in deployment for three structural reasons:"
                    ),
                    BulletListBlock(
                        items=(
                            "Data Sparsity: unstable predictions in low-activity regions.",
                            "Adversarial Noise: bot-driven or media-driven signal inflation.",
                            "Lack of Uncertainty: no confidence estimate attached to decisions.",
                        )
                    ),
                    QuoteBlock(
                        "How can a system estimate credible crisis signals under sparse, noisy, "
                        "and adversarial conditions while producing calibrated confidence for "
                        "human decision-making?"
                    ),
                ),
            ),
            Section(
                title="System Design",
                blocks=(
                    Subsection(
                        title="Overview",
                        blocks=(
                            ParagraphBlock(
                                "The decision score combines inferred crisis likelihood, "
                                "credibility, and suppression for adversarial activity."
                            ),
                            EquationBlock(
                                plain="Credible Signal = chi x P(crisis | s) x (1 - threat)",
                                latex=r"\text{Credible Signal} = \chi \times P(\text{crisis} \mid s) \times (1 - \text{threat})",
                            ),
                        ),
                    ),
                    Subsection(
                        title="Architecture",
                        blocks=(
                            ParagraphBlock("The pipeline is organized into four layers:"),
                            BulletListBlock(
                                items=(
                                    "Data Ingestion",
                                    "Signal Extraction",
                                    "Credibility and Threat Modeling",
                                    "Decision and Governance",
                                )
                            ),
                            ParagraphBlock(
                                "Insert Figure 1 after this subsection. See Figure Production "
                                "Prompts for the exact diagram specification."
                            ),
                        ),
                    ),
                    Subsection(
                        title="Credibility Model",
                        blocks=(
                            EquationBlock(
                                plain="chi = sample_stability x spatial_support x temporal_consistency",
                                latex=r"\chi = \text{sample\_stability} \times \text{spatial\_support} \times \text{temporal\_consistency}",
                            ),
                            EquationBlock(
                                plain="sample_stability = min(1, sqrt(N / k))",
                                latex=r"\text{sample\_stability} = \min(1, \sqrt{N / k})",
                            ),
                            ParagraphBlock(
                                "The credibility term penalizes low-volume regions, rewards "
                                "agreement across neighboring regions, and rejects short-lived "
                                "spikes that diverge from historical behavior."
                            ),
                            ParagraphBlock(
                                "Insert Figure 2 after this subsection. See Figure Production "
                                "Prompts for the exact diagram specification."
                            ),
                        ),
                    ),
                    Subsection(
                        title="Threat Model",
                        blocks=(
                            BulletListBlock(
                                items=(
                                    "Lexical entropy",
                                    "URL density",
                                    "Velocity anomaly",
                                )
                            ),
                            ParagraphBlock(
                                "Threat scoring suppresses coordinated amplification, "
                                "informational spikes, and other non-organic activity before any "
                                "escalation is emitted."
                            ),
                        ),
                    ),
                ),
            ),
            Section(
                title="Implementation Plan",
                blocks=(
                    BulletListBlock(
                        items=(
                            "Phase 1: Data ingestion",
                            "Phase 2: Signal extraction",
                            "Phase 3: Credibility modeling",
                            "Phase 4: Threat filtering",
                            "Phase 5: Decision engine",
                            "Phase 6: Evaluation",
                            "Phase 7: Governance",
                        )
                    ),
                ),
            ),
            Section(
                title="Timeline",
                blocks=(
                    TableBlock(
                        headers=("Week", "Task"),
                        rows=(
                            ("1-2", "Data pipeline"),
                            ("3", "Signal extraction"),
                            ("4", "Credibility modeling"),
                            ("5", "Threat filtering"),
                            ("6", "Decision engine"),
                            ("7", "Evaluation"),
                            ("8", "Governance"),
                        ),
                    ),
                ),
            ),
            Section(
                title="Deliverables",
                blocks=(
                    BulletListBlock(
                        items=(
                            "Credible signal pipeline",
                            "Signal extraction module",
                            "Credibility model (chi)",
                            "Threat filter",
                            "Decision engine",
                            "Evaluation framework",
                            "Audit logging system",
                        )
                    ),
                ),
            ),
            Section(
                title="Evaluation Plan",
                blocks=(
                    LabelValueBlock(
                        label="Primary Metric",
                        value="Precision @ 75% Recall",
                    ),
                    LabelValueBlock(
                        label="Secondary Metrics",
                        value="False Positive Rate, calibration error (ECE), and score stability.",
                    ),
                    ParagraphBlock(
                        "Evaluation uses synthetic sparse-data, bot-amplification, media-spike, "
                        "and organic-crisis scenarios to measure whether the system escalates "
                        "only credible signals."
                    ),
                    ParagraphBlock(
                        "Insert Figure 3 in this section. See Figure Production Prompts for the "
                        "exact scenario diagram specification."
                    ),
                ),
            ),
            Section(
                title="Risks and Mitigation",
                blocks=(
                    BulletListBlock(
                        items=(
                            "Sparse Data: chi penalization and minimum-support gating.",
                            "Adversarial Attacks: entropy and velocity filters.",
                            "False Escalation: mandatory human-in-the-loop review.",
                        )
                    ),
                ),
            ),
            Section(
                title="Why This Project",
                blocks=(
                    ParagraphBlock(
                        "The proposal targets decision reliability rather than classification "
                        "accuracy. It aligns with AI4MH priorities around explainability, "
                        "auditability, and governance while remaining deployable in public "
                        "health monitoring settings."
                    ),
                ),
            ),
            Section(
                title="References",
                blocks=(
                    NumberedListBlock(
                        items=(
                            "Atmakuru et al. (2025). AI-based suicide prevention and prediction: A systematic review. Information Fusion.",
                            "CLPsych shared tasks on suicide-risk detection in social media.",
                            "Moran's I for spatial autocorrelation and regional cluster analysis.",
                            "Bayesian inference and uncertainty-aware decision scoring methods.",
                        )
                    ),
                ),
            ),
            Section(
                title="Figure Production Prompts",
                blocks=(
                    ParagraphBlock(
                        "Limit the proposal to three figures. Do not add UI screenshots, "
                        "generic machine-learning diagrams, or unrelated graphs."
                    ),
                    FigurePromptBlock(
                        title="Figure 1. System Architecture Diagram",
                        placement="Place after System Design -> Architecture.",
                        prompt=(
                            "Draw a clean pipeline diagram:\n\n"
                            "Social Media Data ->\n"
                            "Data Ingestion ->\n"
                            "Signal Extraction ->\n"
                            "Credibility Model (chi) ->\n"
                            "Threat Filter ->\n"
                            "Decision Engine ->\n"
                            "Human Review\n\n"
                            "Use rectangular blocks with arrows.\n"
                            "Highlight chi and threat filter in bold.\n"
                            "Minimal colors, academic style."
                        ),
                    ),
                    FigurePromptBlock(
                        title="Figure 2. Credible Signal Equation Diagram",
                        placement="Place after System Design -> Credibility Model.",
                        prompt=(
                            "Visualize equation:\n\n"
                            "Credible Signal = chi x P(crisis | s) x (1 - threat)\n\n"
                            "Break into 3 components:\n"
                            "1. chi (stability, spatial, temporal)\n"
                            "2. P(crisis | s)\n"
                            "3. threat filter\n\n"
                            "Use labeled blocks feeding into the final score."
                        ),
                    ),
                    FigurePromptBlock(
                        title="Figure 3. Evaluation Scenario Diagram",
                        placement="Place in Evaluation Plan.",
                        prompt=(
                            "Show 4 scenarios:\n\n"
                            "1. Sparse data\n"
                            "2. Bot attack\n"
                            "3. Media spike\n"
                            "4. Organic crisis\n\n"
                            "Use a separate timeline plot for each scenario.\n"
                            "Show expected system response: stable, suppressed, monitored, or detected."
                        ),
                    ),
                ),
            ),
        ),
    )


def escape_tex(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
        "|": r"\textbar{}",
    }
    return "".join(replacements.get(ch, ch) for ch in text)


def render_latex(doc: ProposalDocument) -> str:
    lines: list[str] = [
        r"\documentclass[11pt]{article}",
        r"\usepackage[a4paper,margin=1in]{geometry}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage{amsmath}",
        r"\usepackage{booktabs}",
        r"\usepackage{enumitem}",
        "",
        rf"\title{{{escape_tex(doc.title)}}}",
        rf"\author{{{escape_tex(doc.author)}}}",
        r"\date{}",
        "",
        r"\begin{document}",
        r"\maketitle",
        "",
    ]

    for section in doc.sections:
        lines.extend(render_latex_section(section))

    lines.append(r"\end{document}")
    lines.append("")
    return "\n".join(lines)


def render_latex_section(section: Section) -> list[str]:
    lines = [rf"\section{{{escape_tex(section.title)}}}", ""]
    for block in section.blocks:
        lines.extend(render_latex_block(block, subsection_level=0))
    return lines


def render_latex_block(block: object, subsection_level: int) -> list[str]:
    if isinstance(block, ParagraphBlock):
        return [escape_tex(block.text), ""]
    if isinstance(block, BulletListBlock):
        lines = [r"\begin{itemize}[leftmargin=*]"]
        lines.extend(rf"\item {escape_tex(item)}" for item in block.items)
        lines.extend([r"\end{itemize}", ""])
        return lines
    if isinstance(block, NumberedListBlock):
        lines = [r"\begin{enumerate}[leftmargin=*]"]
        lines.extend(rf"\item {escape_tex(item)}" for item in block.items)
        lines.extend([r"\end{enumerate}", ""])
        return lines
    if isinstance(block, EquationBlock):
        return [r"\begin{equation*}", block.latex, r"\end{equation*}", ""]
    if isinstance(block, QuoteBlock):
        return [r"\begin{quote}", escape_tex(block.text), r"\end{quote}", ""]
    if isinstance(block, TableBlock):
        width_spec = " ".join("l" for _ in block.headers)
        lines = [rf"\begin{{tabular}}{{{width_spec}}}", r"\toprule"]
        lines.append(" & ".join(escape_tex(h) for h in block.headers) + r" \\")
        lines.append(r"\midrule")
        for row in block.rows:
            lines.append(" & ".join(escape_tex(cell) for cell in row) + r" \\")
        lines.extend([r"\bottomrule", r"\end{tabular}", ""])
        return lines
    if isinstance(block, LabelValueBlock):
        return [rf"\textbf{{{escape_tex(block.label)}:}} {escape_tex(block.value)}", ""]
    if isinstance(block, FigurePromptBlock):
        return [
            rf"\subsection*{{{escape_tex(block.title)}}}",
            rf"\textbf{{Placement:}} {escape_tex(block.placement)}",
            "",
            r"\begin{quote}",
            r"\ttfamily",
            escape_tex(block.prompt).replace("\n", r"\\"),
            r"\end{quote}",
            "",
        ]
    if isinstance(block, Subsection):
        lines = [rf"\subsection{{{escape_tex(block.title)}}}", ""]
        for sub_block in block.blocks:
            lines.extend(render_latex_block(sub_block, subsection_level + 1))
        return lines
    raise TypeError(f"Unsupported block type: {type(block)!r}")


def build_docx(doc: ProposalDocument, output_path: Path) -> None:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.start_type = WD_SECTION.NEW_PAGE

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(doc.title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(16)

    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author.add_run(doc.author)
    author_run.font.name = "Times New Roman"
    author_run.font.size = Pt(11)

    for section_data in doc.sections:
        document.add_heading(section_data.title, level=1)
        for block in section_data.blocks:
            render_docx_block(document, block)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def has_docx_style(document, style_name: str) -> bool:
    return any(style.name == style_name for style in document.styles)


def render_docx_block(document, block: object) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    if isinstance(block, ParagraphBlock):
        document.add_paragraph(block.text)
        return
    if isinstance(block, BulletListBlock):
        for item in block.items:
            document.add_paragraph(item, style="List Bullet")
        return
    if isinstance(block, NumberedListBlock):
        for item in block.items:
            document.add_paragraph(item, style="List Number")
        return
    if isinstance(block, EquationBlock):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(block.plain)
        run.font.name = "Courier New"
        return
    if isinstance(block, QuoteBlock):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.right_indent = Inches(0.25)
        paragraph.add_run(block.text).italic = True
        return
    if isinstance(block, TableBlock):
        table = document.add_table(rows=1, cols=len(block.headers))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for index, header in enumerate(block.headers):
            header_cells[index].text = header
        for row in block.rows:
            row_cells = table.add_row().cells
            for index, value in enumerate(row):
                row_cells[index].text = value
        return
    if isinstance(block, LabelValueBlock):
        paragraph = document.add_paragraph()
        label_run = paragraph.add_run(f"{block.label}: ")
        label_run.bold = True
        paragraph.add_run(block.value)
        return
    if isinstance(block, FigurePromptBlock):
        document.add_heading(block.title, level=2)
        placement = document.add_paragraph()
        placement_label = placement.add_run("Placement: ")
        placement_label.bold = True
        placement.add_run(block.placement)
        prompt = document.add_paragraph()
        prompt.style = "Intense Quote" if has_docx_style(document, "Intense Quote") else "Quote"
        prompt.add_run(block.prompt)
        return
    if isinstance(block, Subsection):
        document.add_heading(block.title, level=2)
        for sub_block in block.blocks:
            render_docx_block(document, sub_block)
        return
    raise TypeError(f"Unsupported block type: {type(block)!r}")


def build_pdf(doc: ProposalDocument, output_path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        Paragraph,
        Preformatted,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ProposalTitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        spaceAfter=10,
    )
    author_style = ParagraphStyle(
        "ProposalAuthor",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontName="Helvetica",
        fontSize=11,
        leading=14,
        spaceAfter=18,
    )
    section_style = ParagraphStyle(
        "ProposalSection",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        spaceBefore=10,
        spaceAfter=8,
    )
    subsection_style = ParagraphStyle(
        "ProposalSubsection",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        spaceBefore=8,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ProposalBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        spaceAfter=8,
    )
    label_style = ParagraphStyle(
        "ProposalLabel",
        parent=body_style,
        spaceAfter=8,
    )
    quote_style = ParagraphStyle(
        "ProposalQuote",
        parent=body_style,
        leftIndent=20,
        rightIndent=20,
        italic=True,
    )
    equation_style = ParagraphStyle(
        "ProposalEquation",
        parent=body_style,
        alignment=TA_CENTER,
        fontName="Courier",
        spaceAfter=8,
    )
    prompt_style = ParagraphStyle(
        "ProposalPrompt",
        parent=body_style,
        fontName="Courier",
        fontSize=9.5,
        leading=12,
        leftIndent=12,
        rightIndent=12,
        backColor=colors.whitesmoke,
        borderPadding=8,
        spaceAfter=10,
    )

    story = [
        Paragraph(html.escape(doc.title), title_style),
        Paragraph(html.escape(doc.author), author_style),
    ]

    for section in doc.sections:
        story.append(Paragraph(html.escape(section.title), section_style))
        for block in section.blocks:
            render_pdf_block(
                story=story,
                block=block,
                body_style=body_style,
                subsection_style=subsection_style,
                quote_style=quote_style,
                equation_style=equation_style,
                prompt_style=prompt_style,
                label_style=label_style,
            )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    document.build(story)


def render_pdf_block(
    *,
    story: list[object],
    block: object,
    body_style,
    subsection_style,
    quote_style,
    equation_style,
    prompt_style,
    label_style,
) -> None:
    from reportlab.lib import colors
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, Spacer, Table, TableStyle

    if isinstance(block, ParagraphBlock):
        story.append(Paragraph(html.escape(block.text), body_style))
        return
    if isinstance(block, BulletListBlock):
        items = [
            ListItem(Paragraph(html.escape(item), body_style), leftIndent=10)
            for item in block.items
        ]
        story.append(ListFlowable(items, bulletType="bullet"))
        story.append(Spacer(1, 6))
        return
    if isinstance(block, NumberedListBlock):
        items = [
            ListItem(Paragraph(html.escape(item), body_style), value=index + 1)
            for index, item in enumerate(block.items)
        ]
        story.append(ListFlowable(items, bulletType="1"))
        story.append(Spacer(1, 6))
        return
    if isinstance(block, EquationBlock):
        story.append(Paragraph(html.escape(block.plain), equation_style))
        return
    if isinstance(block, QuoteBlock):
        story.append(Paragraph(html.escape(block.text), quote_style))
        return
    if isinstance(block, TableBlock):
        table_data = [list(block.headers), *[list(row) for row in block.rows]]
        table = Table(table_data, hAlign="LEFT")
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                    ("LEADING", (0, 0), (-1, -1), 12),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 8))
        return
    if isinstance(block, LabelValueBlock):
        story.append(
            Paragraph(
                f"<b>{html.escape(block.label)}:</b> {html.escape(block.value)}",
                label_style,
            )
        )
        return
    if isinstance(block, FigurePromptBlock):
        story.append(Paragraph(html.escape(block.title), subsection_style))
        story.append(
            Paragraph(
                f"<b>Placement:</b> {html.escape(block.placement)}",
                body_style,
            )
        )
        story.append(
            Paragraph(
                html.escape(block.prompt).replace("\n", "<br/>"),
                prompt_style,
            )
        )
        return
    if isinstance(block, Subsection):
        story.append(Paragraph(html.escape(block.title), subsection_style))
        for sub_block in block.blocks:
            render_pdf_block(
                story=story,
                block=sub_block,
                body_style=body_style,
                subsection_style=subsection_style,
                quote_style=quote_style,
                equation_style=equation_style,
                prompt_style=prompt_style,
                label_style=label_style,
            )
        return
    raise TypeError(f"Unsupported block type: {type(block)!r}")


def ensure_parent(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def default_paths(script_path: Path) -> tuple[Path, Path, Path]:
    repo_root = script_path.resolve().parents[1]
    tex_output = repo_root / "docs" / "AI4MH_Proposal.tex"
    pdf_output = repo_root / "output" / "pdf" / "AI4MH.pdf"
    docx_output = repo_root / "output" / "doc" / "AI4MH.docx"
    return tex_output, pdf_output, docx_output


def main() -> int:
    tex_default, pdf_default, docx_default = default_paths(Path(__file__))

    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--tex", type=Path, default=tex_default, help="LaTeX output path")
    parser.add_argument("--pdf", type=Path, default=pdf_default, help="PDF output path")
    parser.add_argument("--docx", type=Path, default=docx_default, help="DOCX output path")
    args = parser.parse_args()

    proposal = build_document()
    ensure_parent(args.tex).write_text(render_latex(proposal), encoding="utf-8")
    build_pdf(proposal, args.pdf)
    build_docx(proposal, args.docx)

    print(f"Wrote LaTeX: {args.tex}")
    print(f"Wrote PDF:   {args.pdf}")
    print(f"Wrote DOCX:  {args.docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
